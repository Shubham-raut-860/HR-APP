"""
Resume file parsing service â€” PDF, DOCX, DOC, and image (OCR) extraction.

â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
SYSTEM DEPENDENCIES (must be installed at OS level â€” pip alone is not enough)
â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

  Ubuntu / Debian:
    sudo apt-get update
    sudo apt-get install -y tesseract-ocr tesseract-ocr-eng poppler-utils

  macOS (Homebrew):
    brew install tesseract poppler

  Windows:
    1. Tesseract installer:  https://github.com/UB-Mannheim/tesseract/wiki
       Default install path: C:\\Program Files\\Tesseract-OCR\\tesseract.exe
       Add to PATH or set pytesseract.pytesseract.tesseract_cmd (see below).
    2. Poppler for Windows:  https://github.com/oschwartz10612/poppler-windows
       Add bin/ folder to PATH.

  Python packages (requirements.txt / pip):
    pdfplumber>=0.11.0
    python-docx>=1.1.0
    docx2txt>=0.8
    Pillow>=10.3.0
    pytesseract>=0.3.13
    pdf2image>=1.17.0

â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
Supported input types
â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
  .pdf          â†’ pdfplumber first; auto-falls back to Tesseract OCR
                  when the PDF is scanned / image-only.
  .docx         â†’ python-docx (paragraphs + tables + headers/footers)
  .doc          â†’ docx2txt (legacy Word format)
  .png .jpg     â†’ Tesseract OCR via pytesseract
  .jpeg .webp   â†’ Tesseract OCR via pytesseract
  .tiff .tif    â†’ Tesseract OCR via pytesseract (best for scanned docs)
  .bmp .gif     â†’ Tesseract OCR via pytesseract

OCR pipeline
â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
  1. Normalise image mode (RGB â†’ grayscale).
  2. Upscale to â‰¥ 300 DPI equivalent if the image is small.
  3. Apply a mild sharpen to recover edge definition lost in scanning.
  4. Run Tesseract PSM 3 (fully automatic page segmentation).
  5. If result is sparse (<60 chars), retry with PSM 6 (uniform block)
     and keep whichever produced more content.
"""
import io
import asyncio
import logging
from pathlib import Path

from fastapi import UploadFile, HTTPException, status
import pdfplumber
from docx import Document
import docx2txt
from PIL import Image, ImageFilter
import pytesseract
import pypdfium2 as pdfium

from app.config import settings
from app.services import encryption_service

logger = logging.getLogger(__name__)

# â”€â”€â”€ Extension sets â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".tiff", ".tif", ".bmp", ".gif"}

# BUG-12 FIX: ALL_ALLOWED_EXTENSIONS was a hardcoded set that could silently
# diverge from config.ALLOWED_RESUME_EXTENSIONS. Upload endpoints validated
# against config; file_service routed against its own set. Adding ".rtf" to
# config would pass the upload gate but crash here with an unhandled error.
#
# Fix: derive from config as the single source of truth.
# The hardcoded IMAGE_EXTENSIONS subtree is kept for the OCR routing decision
# (we must know which extensions need OCR regardless of config).
ALL_ALLOWED_EXTENSIONS: set[str] = set(settings.allowed_extensions_list)

# Startup assertion: warn if config lists an extension file_service can't route.
# Extend _EXTRACTABLE when a new format is supported (parser added below).
_EXTRACTABLE = {".pdf", ".docx", ".doc", ".txt"} | IMAGE_EXTENSIONS
_unroutable = ALL_ALLOWED_EXTENSIONS - _EXTRACTABLE
if _unroutable:
    import warnings as _w
    _w.warn(
        f"BUG-12: config.ALLOWED_RESUME_EXTENSIONS contains extensions that "
        f"file_service cannot extract text from: {sorted(_unroutable)}. "
        f"Either add a parser for them or remove them from the config.",
        stacklevel=2,
    )

# pdfplumber results shorter than this trigger OCR fallback
_PDF_TEXT_THRESHOLD = 600
_PDF_PAGE_TEXT_THRESHOLD = 80
_OCR_MIN_SIDE_PX = 1400
_OCR_DPI = 300
_MAX_OCR_PAGES = 20

# â”€â”€â”€ Tesseract availability check â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# BUG FIX: Previously there was no startup check. If tesseract-ocr is not
# installed at the OS level, every OCR call would raise a confusing FileNotFoundError
# or "tesseract is not installed" error with no guidance. We now detect this once
# at import time and surface a clear, actionable message.

_TESSERACT_AVAILABLE = False


def _configure_tesseract_path() -> None:
    """
    On Windows, Tesseract is installed to a non-PATH location by default.
    The UB-Mannheim installer puts it at:
      C:\\Program Files\\Tesseract-OCR\\tesseract.exe
    Many users install it correctly but never add it to PATH, so
    pytesseract.get_tesseract_version() fails even though the binary exists.

    Probes common Windows install locations and sets tesseract_cmd if found.
    No-op on Linux/macOS.
    """
    import sys
    import os
    if sys.platform != "win32":
        return
    candidates = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        os.path.expanduser(r"~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"),
        os.path.expanduser(r"~\AppData\Local\Tesseract-OCR\tesseract.exe"),
    ]
    for path in candidates:
        if os.path.isfile(path):
            pytesseract.pytesseract.tesseract_cmd = path
            logger.info("Tesseract found at %s â€” configured automatically.", path)
            return


_configure_tesseract_path()

try:
    pytesseract.get_tesseract_version()
    _TESSERACT_AVAILABLE = True
except Exception as _tess_err:
    logger.warning(
        "Tesseract OCR binary not found (%s). "
        "OCR will be unavailable until you install it:\n"
        "  Ubuntu/Debian: sudo apt-get install -y tesseract-ocr tesseract-ocr-eng\n"
        "  macOS:         brew install tesseract\n"
        "  Windows:       https://github.com/UB-Mannheim/tesseract/wiki\n"
        "                 Default path after install: C:\\Program Files\\Tesseract-OCR\\tesseract.exe\n"
        "                 Tick 'Add to PATH' during install, or restart the server after installing.",
        _tess_err,
    )


# â”€â”€â”€ Validation â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def _validate_extension(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    # BUG-12 FIX: use ALL_ALLOWED_EXTENSIONS (derived from config) as the single
    # source of truth. No longer need the union with a separate hardcoded set.
    if ext not in ALL_ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"File type '{ext}' is not supported. "
                f"Accepted formats: {', '.join(sorted(ALL_ALLOWED_EXTENSIONS))}"
            ),
        )
    return ext


def _validate_size(content: bytes) -> None:
    size_mb = len(content) / (1024 * 1024)
    if size_mb > settings.MAX_FILE_SIZE_MB:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File exceeds the {settings.MAX_FILE_SIZE_MB} MB size limit",
        )


def _looks_like_internal_ciphertext(content: bytes) -> bool:
    """
    Detect encrypted-at-rest blobs generated by encryption_service.
    Supports legacy Fernet and chunked AES-GCM payloads.
    """
    if len(content) < 32:
        return False
    return content.startswith(b"gAAAAA") or content.startswith(b"HRAPPA2\x00")


def _maybe_decrypt_uploaded_content(filename: str, ext: str, content: bytes) -> bytes:
    """
    If users accidentally upload files from uploads/resumes (already encrypted),
    decrypt them before parsing. If decryption fails, raise a clear actionable
    error instead of noisy PDF/OCR parser failures.
    """
    if ext == ".txt" or not _looks_like_internal_ciphertext(content):
        return content

    decrypted = encryption_service.try_decrypt_file(content)
    if decrypted is not None:
        logger.warning(
            "Detected encrypted upload artifact for %s; auto-decrypted before parsing.",
            filename,
        )
        return decrypted

    raise HTTPException(
        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
        detail=(
            "Uploaded file appears to be an encrypted internal storage copy "
            "(from uploads/resumes), not an original resume document. "
            "Please upload the original source file."
        ),
    )


# â”€â”€â”€ OCR helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def _preprocess_for_ocr(img: Image.Image) -> Image.Image:
    """
    Prepare a PIL image for Tesseract:
      â€¢ Normalise to RGB then grayscale.
      â€¢ Upscale tiny images so Tesseract gets enough pixels.
      â€¢ Sharpen to recover edges blurred by scanning/compression.
    """
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    min_side = min(img.width, img.height)
    if min_side < _OCR_MIN_SIDE_PX:
        scale = _OCR_MIN_SIDE_PX / min_side
        img = img.resize((int(img.width * scale), int(img.height * scale)), Image.LANCZOS)

    img = img.convert("L")
    img = img.filter(ImageFilter.SHARPEN)
    return img


def _run_tesseract(img: Image.Image) -> str:
    """
    OCR one PIL image.  Tries PSM 3 first; falls back to PSM 6 when the
    first pass yields very little text (common with dense resume layouts).
    """
    if not _TESSERACT_AVAILABLE:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=(
                "OCR is unavailable: Tesseract is not installed on this server. "
                "Install it with: sudo apt-get install -y tesseract-ocr tesseract-ocr-eng"
            ),
        )

    preprocessed = _preprocess_for_ocr(img)
    # PSM 1 (Automatic page segmentation with OSD) is better for standard resume layouts
    text = pytesseract.image_to_string(preprocessed, lang="eng", config="--oem 3 --psm 1")

    # Fallback to PSM 3 and PSM 6 for extremely dense multi-column or block layouts
    if len(text.strip()) < 500:
        alt = pytesseract.image_to_string(preprocessed, lang="eng", config="--oem 3 --psm 3")
        if len(alt.strip()) < 500:
            alt = pytesseract.image_to_string(preprocessed, lang="eng", config="--oem 3 --psm 6")
        if len(alt.strip()) > len(text.strip()):
            text = alt

    return text.strip()


def _ocr_image_bytes(content: bytes, ext: str) -> str:
    """OCR a raw image file.

    Animated GIF: tries each frame and returns the one that yields the most
    text. Frame 0 is often blank so we keep scanning until we find content.
    """
    try:
        img = Image.open(io.BytesIO(content))
        n_frames = getattr(img, "n_frames", 1)
        if n_frames <= 1:
            return _run_tesseract(img)

        best = ""
        for frame_idx in range(n_frames):
            try:
                img.seek(frame_idx)
                candidate = _run_tesseract(img.copy())
                if len(candidate) > len(best):
                    best = candidate
                    if len(best) > 300:
                        break
            except Exception:
                continue
        return best
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"OCR failed for image file ({ext}): {exc}",
        )


def _ocr_pdf_pages(content: bytes, page_indices: list[int] | None = None) -> tuple[dict[int, str], bool]:
    """
    OCR selected PDF pages (or all pages when page_indices is None).
    Returns ({page_index: text}, truncated_flag).
    """
    try:
        pdf = pdfium.PdfDocument(content)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Could not open PDF for OCR rendering: {exc}",
        )

    total_pages = len(pdf)
    if page_indices is None:
        selected_pages = list(range(total_pages))
    else:
        selected_pages = sorted({i for i in page_indices if 0 <= i < total_pages})

    is_truncated = len(selected_pages) > _MAX_OCR_PAGES
    if is_truncated:
        logger.warning(
            "OCR requested for %d pages â€” capped at %d.",
            len(selected_pages), _MAX_OCR_PAGES,
        )
        selected_pages = selected_pages[:_MAX_OCR_PAGES]

    page_texts: dict[int, str] = {}
    scale = _OCR_DPI / 72  # pypdfium2 default is 72 DPI; scale factor adjusts
    for page_idx in selected_pages:
        try:
            page = pdf[page_idx]
            bitmap = page.render(scale=scale)
            pil_image = bitmap.to_pil()
            page_texts[page_idx] = _run_tesseract(pil_image)
        except HTTPException:
            raise
        except Exception as exc:
            logger.warning("OCR render failed for page %d: %s", page_idx + 1, exc)
            continue

    pdf.close()
    return page_texts, is_truncated


def _ocr_scanned_pdf(content: bytes) -> str:
    """
    OCR all pages of a scanned PDF.
    """
    page_texts, is_truncated = _ocr_pdf_pages(content, page_indices=None)
    text_result = "\n\n".join(
        page_texts[idx] for idx in sorted(page_texts) if page_texts[idx]
    )
    if is_truncated:
        text_result += "\n\n[SYSTEM: OCR_TRUNCATED]"
    return text_result


# â”€â”€â”€ Per-format text extractors â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€


def _text_looks_valid(text: str) -> bool:
    """Return True when extracted text looks like real readable content.
    Catches encoding garbage (e.g. ÃŒÃŽÃÃŽÃ±Ãª) from cheaply-digitised PDFs.
    At least 80% of characters should be printable."""
    if not text or len(text) < 10:
        return False
    # FIX Finding 19: Use isprintable() instead of hardcoded ASCII range (32-127) 
    # to avoid rejecting valid non-English resumes (UTF-8).
    printable_chars = sum(1 for c in text if c.isprintable() or c.isspace())
    return (printable_chars / len(text)) >= 0.80


def extract_text_from_pdf(content: bytes) -> str:
    """
    PDF extraction with automatic scanned-PDF fallback.

    Step 1 â€” pdfplumber: fast, zero-loss extraction for digital PDFs.
    Step 2 â€” pdf2image + Tesseract: triggered when the PDF appears to be
              scanned (very little selectable text found by pdfplumber).
    """
    parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                # layout=True preserves column positions so two-column resumes
                # are read column-by-column instead of interleaved row-by-row.
                t = page.extract_text(layout=True, x_tolerance=5, y_tolerance=5)
                parts.append((t or "").strip())
    except Exception as exc:
        logger.warning("pdfplumber extraction failed (%s); will attempt OCR.", exc)

    digital_text = "\n\n".join(t for t in parts if t).strip()
    if not parts:
        return _ocr_scanned_pdf(content)

    # PDF-first fast path: only mark a page weak when it has no text or the text
    # clearly looks corrupted. Short but valid text should not trigger OCR by itself.
    weak_pages = [
        idx for idx, text in enumerate(parts)
        if (not text) or (not _text_looks_valid(text))
    ]

    # Happy path: every page produced valid text and whole-doc signal is strong.
    if not weak_pages and len(digital_text) >= _PDF_TEXT_THRESHOLD and _text_looks_valid(digital_text):
        return digital_text

    # Image-only / fully weak PDF: OCR whole document as backup.
    if len(weak_pages) == len(parts):
        logger.info(
            "pdfplumber found no reliable text (%d chars, valid=%s) â€” running full OCR fallback.",
            len(digital_text), _text_looks_valid(digital_text),
        )
        try:
            ocr_text = _ocr_scanned_pdf(content)
        except HTTPException as exc:
            if digital_text:
                logger.warning("OCR fallback failed (%s); returning digital text instead.", exc.detail)
                return digital_text
            raise
        return ocr_text if len(ocr_text) > len(digital_text) else digital_text

    # Mixed PDF: for mostly-digital PDFs, skip OCR entirely to keep bulk latency low.
    weak_ratio = len(weak_pages) / max(1, len(parts))
    if weak_ratio <= 0.20 and len(digital_text) >= (_PDF_TEXT_THRESHOLD // 2):
        return digital_text

    # Otherwise, keep digital text where present and OCR only weak pages.
    logger.info(
        "pdfplumber extracted reliable text on %d/%d pages; OCR fallback on %d weak pages.",
        len(parts) - len(weak_pages), len(parts), len(weak_pages),
    )
    try:
        ocr_by_page, is_truncated = _ocr_pdf_pages(content, page_indices=weak_pages)
    except HTTPException as exc:
        if digital_text:
            logger.warning("OCR fallback failed (%s); returning digital text instead.", exc.detail)
            return digital_text
        raise

    merged_pages: list[str] = []
    for idx, text in enumerate(parts):
        ocr_text = (ocr_by_page.get(idx) or "").strip()
        base_text = text.strip()
        if ocr_text and len(ocr_text) > len(base_text):
            merged_pages.append(ocr_text)
        else:
            merged_pages.append(base_text)

    merged_text = "\n\n".join(t for t in merged_pages if t).strip()
    if is_truncated and merged_text:
        merged_text += "\n\n[SYSTEM: OCR_TRUNCATED]"
    return merged_text or digital_text


def extract_text_from_docx(content: bytes) -> str:
    """
    DOCX extraction: body paragraphs + tables + section headers/footers.
    Table cells are joined with spaces so skills grids and education tables
    read as coherent lines rather than single-char fragments.
    """
    try:
        doc = Document(io.BytesIO(content))
        parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = "  ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)

        for section in doc.sections:
            for hf in (section.header, section.footer):
                if hf:
                    for para in hf.paragraphs:
                        if para.text.strip():
                            parts.append(para.text)

        # FIX 2: text boxes (w:txbxContent) â€” widely used in Canva/Google Docs templates
        # for sidebars, headers, skill grids. python-docx paragraphs/tables never visits them.
        try:
            from docx.oxml.ns import qn as _qn
            seen: set[str] = set(parts)
            for txbx in doc.element.body.iter(_qn("w:txbxContent")):
                for p_elem in txbx.iter(_qn("w:p")):
                    text = "".join(r.text for r in p_elem.iter(_qn("w:t")) if r.text)
                    if text.strip() and text not in seen:
                        parts.append(text)
                        seen.add(text)
        except Exception:
            pass  # non-critical: body text already captured above

        return "\n".join(parts)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Failed to parse DOCX: {exc}",
        )


def extract_text_from_doc(content: bytes) -> str:
    """Legacy .doc extraction via docx2txt."""
    try:
        text = docx2txt.process(io.BytesIO(content))
        return (text or "").strip()
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Failed to parse DOC: {exc}",
        )


def extract_text_from_txt(content: bytes) -> str:
    """Plain-text extraction with conservative encoding fallbacks."""
    encodings = ("utf-8-sig", "utf-16", "cp1252", "latin-1")
    for encoding in encodings:
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise HTTPException(
        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
        detail="Failed to decode TXT file content.",
    )


# â”€â”€â”€ Public entry points â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

async def extract_text(file: UploadFile) -> tuple[str, bytes]:
    """
    Validate, read, and extract plain text from an uploaded resume file.

    Returns:
        (extracted_text: str, raw_bytes: bytes)
    """
    ext = _validate_extension(file.filename or "")
    content = await file.read()
    _validate_size(content)
    content = _maybe_decrypt_uploaded_content(file.filename or "", ext, content)
    text = await _dispatch_extraction(ext, content)

    if not text.strip():
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                "Could not extract any text from the uploaded file. "
                "For scanned documents, ensure the scan resolution is at least 300 DPI "
                "and the page is not rotated more than 45Â°."
            ),
        )

    return text, content


async def extract_text_from_bytes(filename: str, content: bytes) -> str:
    """
    Extract plain text from already-read bytes.
    Used by bulk upload so bytes are pre-read once, not inside concurrent tasks.
    """
    ext = _validate_extension(filename)
    _validate_size(content)
    content = _maybe_decrypt_uploaded_content(filename, ext, content)
    text = await _dispatch_extraction(ext, content)

    if not text.strip():
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not extract any text from the uploaded file.",
        )

    return text


async def _dispatch_extraction(ext: str, content: bytes) -> str:
    """Central dispatch â€” avoids duplicating the if/elif chain in two places."""
    if ext == ".pdf":
        return await asyncio.to_thread(extract_text_from_pdf, content)
    elif ext == ".docx":
        return await asyncio.to_thread(extract_text_from_docx, content)
    elif ext == ".doc":
        return await asyncio.to_thread(extract_text_from_doc, content)
    elif ext == ".txt":
        return await asyncio.to_thread(extract_text_from_txt, content)
    elif ext in IMAGE_EXTENSIONS:
        return await asyncio.to_thread(_ocr_image_bytes, content, ext)
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type: {ext}",
        )


# â”€â”€â”€ File storage â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

async def save_file(content: bytes, filename: str, subfolder: str = "resumes") -> str:
    """Encrypt and persist a file, returning the on-disk path.

    All disk operations are dispatched via asyncio.to_thread so a slow/
    network-backed UPLOAD_DIR never blocks the ASGI event loop.
    """
    safe_name = (
        "".join(c if (c.isalnum() or c in "._-") else "_" for c in filename).strip("_")
        or "resume"
    )
    def _prepare_and_write(upload_dir: str, sname: str, data: bytes) -> str:
        upload_path = Path(upload_dir) / subfolder
        upload_path.mkdir(parents=True, exist_ok=True)

        import os
        dest = upload_path / sname
        counter = 1
        while True:
            try:
                fd = os.open(dest, os.O_CREAT | os.O_EXCL | os.O_WRONLY, 0o644)
                with os.fdopen(fd, "wb"):
                    pass
                try:
                    encryption_service.encrypt_file_to_path(data, str(dest))
                except Exception:
                    try:
                        os.unlink(dest)
                    except OSError:
                        pass
                    raise
                break
            except FileExistsError:
                stem, suffix = Path(sname).stem, Path(sname).suffix
                dest = upload_path / f"{stem}_{counter}{suffix}"
                counter += 1
        return str(dest)

    return await asyncio.to_thread(_prepare_and_write, settings.UPLOAD_DIR, safe_name, content)
