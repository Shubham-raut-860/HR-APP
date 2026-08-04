import os
import time
import json
import random
import signal
import shutil
import subprocess
from pathlib import Path
from datetime import datetime

import requests
from docx import Document

ROOT = Path(r"D:\Shubham\HR APP\HR APP BACKEND")
DATA_DIR = ROOT / "test_docs" / "perf_bulk_500"
REPORT_PATH = ROOT / "perf_native_vs_harness_report.json"

SIZES = [100, 250, 500]
PORTS = {"native": 8100, "harness": 8101}
POLL_INTERVAL = 2.0
HEALTH_TIMEOUT = 120
RUN_TIMEOUT = 3600  # per run
PASSWORD = "Qa!Pass2026A"

JD_BASE = {
    "title": "Perf Backend Engineer",
    "role": "Backend Engineer",
    "location": "Remote",
    "employment_type": "Full-time",
    "experience_min": 2,
    "experience_max": 7,
    "must_have_skills": ["Python", "FastAPI", "SQL", "PostgreSQL", "Docker", "REST API"],
    "good_to_have_skills": ["AWS", "Redis", "Kubernetes", "CI/CD"],
    "description": "Performance benchmark job for bulk resume ingestion.",
    "resume_weight": 50,
    "quiz_weight": 50,
    "pass_threshold": 60,
}


def ensure_dataset():
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    existing = sorted(DATA_DIR.glob("resume_*.docx"))
    if len(existing) >= 500:
        return existing[:500]

    # Rebuild cleanly for deterministic runs
    for p in DATA_DIR.glob("resume_*.docx"):
        p.unlink(missing_ok=True)

    skill_sets = [
        ["Python", "FastAPI", "SQL", "PostgreSQL", "Docker", "REST API", "AWS"],
        ["Python", "Django", "PostgreSQL", "Redis", "Docker"],
        ["Java", "Spring", "MySQL"],
        ["Node.js", "Express", "MongoDB"],
        ["Python", "FastAPI", "Kubernetes", "CI/CD", "SQL"],
    ]

    for i in range(1, 501):
        name = f"Perf Candidate {i:03d}"
        email = f"perf.candidate.{i:03d}@example.com"
        yrs = (i % 9) + 1
        skills = skill_sets[i % len(skill_sets)]
        doc = Document()
        doc.add_heading(name, level=1)
        doc.add_paragraph(f"Email: {email}")
        doc.add_paragraph(f"Phone: +91-90000{i:05d}")
        doc.add_paragraph("Location: India")
        doc.add_paragraph(f"Experience: {yrs} years")
        doc.add_paragraph("Skills: " + ", ".join(skills))
        doc.add_paragraph("Summary: Backend engineer profile for load benchmark.")
        doc.add_paragraph("Projects: Built APIs, optimized SQL queries, containerized deployments.")
        out = DATA_DIR / f"resume_{i:03d}.docx"
        doc.save(out)

    return sorted(DATA_DIR.glob("resume_*.docx"))[:500]


def start_server(mode: str):
    port = PORTS[mode]
    env = os.environ.copy()
    env["BULK_USE_HARNESS_PIPELINE"] = "true" if mode == "harness" else "false"
    env["AI_SCORE_CACHE_ENABLED"] = "false"
    env["BULK_FAST_MODE"] = "true"

    out_log = ROOT / f"perf_{mode}_server.out.log"
    err_log = ROOT / f"perf_{mode}_server.err.log"
    out_f = open(out_log, "w", encoding="utf-8")
    err_f = open(err_log, "w", encoding="utf-8")

    proc = subprocess.Popen(
        [str(ROOT / ".venv" / "Scripts" / "python.exe"), "-m", "uvicorn", "app.main:app", "--host", "127.0.0.1", "--port", str(port)],
        cwd=str(ROOT),
        env=env,
        stdout=out_f,
        stderr=err_f,
    )

    return proc, out_f, err_f, port


def stop_server(proc, out_f, err_f):
    try:
        if proc.poll() is None:
            proc.terminate()
            try:
                proc.wait(timeout=20)
            except subprocess.TimeoutExpired:
                proc.kill()
                proc.wait(timeout=10)
    finally:
        out_f.close()
        err_f.close()


def wait_health(base_url: str, timeout_s=HEALTH_TIMEOUT):
    t0 = time.time()
    while time.time() - t0 < timeout_s:
        try:
            r = requests.get(f"{base_url}/health", timeout=5)
            if r.status_code == 200:
                return True
        except Exception:
            pass
        time.sleep(1)
    return False


def register_and_login(base_url: str, mode: str):
    stamp = int(time.time())
    email = f"perf.{mode}.{stamp}.{random.randint(1000,9999)}@example.com"
    reg_payload = {
        "full_name": f"Perf {mode.title()} User",
        "email": email,
        "password": PASSWORD,
        "role": "hr",
    }
    reg = requests.post(f"{base_url}/auth/register", json=reg_payload, timeout=30)
    if reg.status_code not in (200, 201):
        raise RuntimeError(f"register failed: {reg.status_code} {reg.text}")

    login = requests.post(
        f"{base_url}/auth/login",
        json={"email": email, "password": PASSWORD},
        timeout=30,
    )
    if login.status_code != 200:
        raise RuntimeError(f"login failed: {login.status_code} {login.text}")
    data = login.json()
    return data["access_token"], email


def create_job(base_url: str, token: str, mode: str, n: int):
    payload = dict(JD_BASE)
    payload["title"] = f"Perf {mode} job {n} {int(time.time())}"
    payload["role"] = f"Perf {mode} role"
    r = requests.post(
        f"{base_url}/jd/",
        json=payload,
        headers={"Authorization": f"Bearer {token}"},
        timeout=120,
    )
    if r.status_code not in (200, 201):
        raise RuntimeError(f"create job failed: {r.status_code} {r.text}")
    return r.json()["id"], r.json().get("title")


def run_bulk_async(base_url: str, token: str, job_id: str, files: list[Path], mode: str, n: int):
    multipart = []
    handles = []
    try:
        for idx, fp in enumerate(files, start=1):
            h = open(fp, "rb")
            handles.append(h)
            multipart.append(("files", (fp.name, h, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")))
            multipart.append(("file_ids", (None, f"{mode}_{n}_{idx:04d}")))
        multipart.append(("job_id", (None, job_id)))

        t_submit_start = time.time()
        submit = requests.post(
            f"{base_url}/resumes/upload-bulk-async",
            files=multipart,
            headers={"Authorization": f"Bearer {token}"},
            timeout=600,
        )
        t_submit_end = time.time()

        if submit.status_code not in (200, 201, 202):
            raise RuntimeError(f"submit failed: {submit.status_code} {submit.text[:1000]}")

        sjson = submit.json()
        run_id = sjson["job_id"]

        poll_url = f"{base_url}/resumes/upload-bulk-async/{run_id}"
        t0 = time.time()
        progress_points = []
        first_processed_ts = None
        last_processed = 0
        max_stall = 0.0
        last_change_ts = t0

        while True:
            pr = requests.get(poll_url, headers={"Authorization": f"Bearer {token}"}, timeout=30)
            if pr.status_code != 200:
                raise RuntimeError(f"poll failed: {pr.status_code} {pr.text[:500]}")
            pdata = pr.json()
            status = pdata.get("status")
            prog = pdata.get("progress") or {}
            processed = int(prog.get("processed") or 0)
            total = int(prog.get("total") or len(files))
            now = time.time()

            progress_points.append({
                "t": round(now - t0, 2),
                "status": status,
                "processed": processed,
                "total": total,
                "success_count": int(prog.get("success_count") or 0),
                "failed_count": int(prog.get("failed_count") or 0),
                "duplicate_count": int(prog.get("duplicate_count") or 0),
            })

            if processed > 0 and first_processed_ts is None:
                first_processed_ts = now

            if processed > last_processed:
                stall = now - last_change_ts
                if stall > max_stall:
                    max_stall = stall
                last_change_ts = now
                last_processed = processed

            if status in ("completed", "failed"):
                done = pdata
                break

            if now - t0 > RUN_TIMEOUT:
                raise RuntimeError(f"run timeout > {RUN_TIMEOUT}s")

            time.sleep(POLL_INTERVAL)

        t1 = time.time()
        result_summary = ((done.get("result") or {}).get("summary") or {})
        total_s = t1 - t0
        success_count = int(result_summary.get("success_count") or 0)
        failed_count = int(result_summary.get("failed_count") or 0)
        dup_count = int(result_summary.get("duplicate_count") or 0)

        return {
            "mode": mode,
            "size": n,
            "submit_latency_s": round(t_submit_end - t_submit_start, 2),
            "async_total_s": round(total_s, 2),
            "time_to_first_processed_s": round((first_processed_ts - t0), 2) if first_processed_ts else None,
            "max_stall_between_progress_s": round(max_stall, 2),
            "throughput_files_per_s": round((success_count / total_s), 3) if total_s > 0 else None,
            "status": done.get("status"),
            "orchestrator": result_summary.get("orchestrator") or done.get("orchestrator"),
            "accepted_count": int(done.get("accepted_count") or n),
            "success_count": success_count,
            "failed_count": failed_count,
            "duplicate_count": dup_count,
            "http_status": (done.get("result") or {}).get("http_status"),
            "progress_tail": progress_points[-10:],
        }
    finally:
        for h in handles:
            try:
                h.close()
            except Exception:
                pass


def benchmark_mode(mode: str, all_files: list[Path]):
    proc, out_f, err_f, port = start_server(mode)
    base = f"http://127.0.0.1:{port}"
    mode_result = {"mode": mode, "runs": [], "server_port": port}

    try:
        if not wait_health(base):
            raise RuntimeError(f"{mode} server failed health check on {base}")

        token, email = register_and_login(base, mode)
        mode_result["recruiter_email"] = email

        for n in SIZES:
            job_id, job_title = create_job(base, token, mode, n)
            subset = all_files[:n]
            run = run_bulk_async(base, token, job_id, subset, mode, n)
            run["job_id"] = job_id
            run["job_title"] = job_title
            mode_result["runs"].append(run)

    finally:
        stop_server(proc, out_f, err_f)

    return mode_result


def main():
    files = ensure_dataset()
    stamp = datetime.utcnow().isoformat()

    report = {
        "generated_at_utc": stamp,
        "dataset_dir": str(DATA_DIR),
        "dataset_count": len(files),
        "sizes": SIZES,
        "results": [],
        "summary": {},
    }

    for mode in ("native", "harness"):
        print(f"\n=== Benchmark mode: {mode} ===")
        mr = benchmark_mode(mode, files)
        report["results"].append(mr)
        for r in mr["runs"]:
            print(
                f"{mode:7s} n={r['size']:3d} total={r['async_total_s']:7.2f}s "
                f"submit={r['submit_latency_s']:6.2f}s ttfp={r['time_to_first_processed_s']} "
                f"thr={r['throughput_files_per_s']}/s ok={r['success_count']} fail={r['failed_count']} dup={r['duplicate_count']}"
            )

    # Build paired summary
    by_mode = {m["mode"]: {r["size"]: r for r in m["runs"]} for m in report["results"]}
    paired = []
    for n in SIZES:
        nat = by_mode.get("native", {}).get(n)
        har = by_mode.get("harness", {}).get(n)
        if not nat or not har:
            continue
        paired.append({
            "size": n,
            "native_total_s": nat["async_total_s"],
            "harness_total_s": har["async_total_s"],
            "delta_s": round(har["async_total_s"] - nat["async_total_s"], 2),
            "native_thr": nat["throughput_files_per_s"],
            "harness_thr": har["throughput_files_per_s"],
            "speedup_native_over_harness": round((har["async_total_s"] / nat["async_total_s"]), 3) if nat["async_total_s"] else None,
        })
    report["summary"]["paired"] = paired

    REPORT_PATH.write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(f"\nReport saved: {REPORT_PATH}")


if __name__ == "__main__":
    main()
